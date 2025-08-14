import fitz  # PyMuPDF
import google.generativeai as genai
from PIL import Image
import io
import json
import os
from typing import List, Dict, Any
import time
from dotenv import load_dotenv
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import re

##############
##%

class ICMemoAnalyzer:
    def __init__(self, api_key: str, model_name: str = "gemini-2.5-flash-lite"):
        """
        Initialize the IC Memo Analyzer with Gemini Vision

        Args:
            api_key: Your Google AI API key
            model_name: Gemini model to use (gemini-1.5-pro recommended for vision)
        """
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel(model_name)

    def pdf_to_images(self, pdf_path: str, dpi: int = 300) -> List[Image.Image]:
        """
        Convert PDF pages to high-resolution images

        Args:
            pdf_path: Path to the PDF file
            dpi: Resolution for image conversion (higher = better quality)

        Returns:
            List of PIL Images, one per page
        """
        pdf_document = fitz.open(pdf_path)
        images = []

        # Calculate scaling factor for desired DPI
        # PyMuPDF default is 72 DPI
        scale = dpi / 72.0
        matrix = fitz.Matrix(scale, scale)

        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]

            # Convert page to image
            pix = page.get_pixmap(matrix=matrix)
            img_data = pix.tobytes("png")

            # Convert to PIL Image
            image = Image.open(io.BytesIO(img_data))
            images.append(image)

            print(f"Converted page {page_num + 1}/{pdf_document.page_count}")

        pdf_document.close()
        return images

    def analyze_page(self, image: Image.Image, page_num: int) -> Dict[str, Any]:
        """
        Analyze a single page using Gemini Vision

        Args:
            image: PIL Image of the page
            page_num: Page number for reference

        Returns:
            Dictionary containing extracted information
        """
        prompt = """
        # Persona
        You are a financial analyst working at Strada Partners, a  a small-cap private equity firm based in Belgium.
        # Task
        Your task is to analyze IC (Investment Committee) memos and extract key information.
        Be thorough and precise with numbers, percentages, and specific details shown in the document.
        # Output
        Output a text with all the relevant information, such that it can be further processed
        """

        try:
            response = self.model.generate_content([prompt, image])

            return {
                "page_number": page_num + 1,
                "content": response.text,
                "status": "success"
            }

        except Exception as e:
            return {
                "page_number": page_num + 1,
                "content": f"Error analyzing page: {str(e)}",
                "status": "error"
            }

    def analyze_pdf(self, pdf_path: str, output_file: str = None, delay: float = 1.0) -> List[Dict[str, Any]]:
        """
        Analyze entire PDF document

        Args:
            pdf_path: Path to the PDF file
            output_file: Optional path to save results as JSON
            delay: Delay between API calls to avoid rate limits

        Returns:
            List of analysis results, one per page
        """
        print(f"Starting analysis of {pdf_path}")

        # Convert PDF to images
        images = self.pdf_to_images(pdf_path)

        # Analyze each page
        results = []
        for i, image in enumerate(images):
            print(f"Analyzing page {i + 1}/{len(images)}...")

            result = self.analyze_page(image, i)
            results.append(result)

            # Add delay to avoid rate limiting
            if i < len(images) - 1:  # Don't delay after the last page
                time.sleep(delay)

        # Save results if output file specified
        if output_file:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
            print(f"Results saved to {output_file}")

        return results

    def create_IC_minutes(self,results, model_name="gemini-2.5-pro"):
        '''
        this function creates IC minutes from the results of the analysis
        :param results:
        :param model_name:
        :return:
        '''
        text = ""
        for i in range(len(results)):
            text += "results from page " + str(i + 1) + "\n" + results[i]["content"] + "\n\n"


        Prompt_Template = '''
        ## Persona
        You are an expert investment analyst tasked with summarizing a detailed investment memo into a formal and concise Investment Committee (IC) Minutes document. 
        You work at Strada Partners, a small-cap private equity firm based in Belgium. You need to draft IC minutes for the investment memo below:
        ## Task
        Your task is to analyze IC (Investment Committee) memos and extract key information.
        Be thorough and precise with numbers, percentages, and specific details shown in the document.
        ## Input
        The input is a PDF file containing an investment memo. The PDF file is provided as a single file.
        ## Output
        Your output must be structured into the following four sections, in this specific order:
        1. Context
        2. Investment Overview
        3. Ask
        4. Conclusion
        Use the provided text, which is an extraction from an investment memo on "Project Calc," to generate the IC minutes.

        ## Detailed description
        1. Context
            - Introduce the Company being acquired
            - Condense the most relevant information from the memo
            - Provide a brief summary of the company's business model, industry, and positioning
        2. Investment Overview
            - Provide a detailed overview of the investment, including the company's name, the type of investment, the investment objective, and the investment strategy
            - Provide a detailed description of the investment objective and strategy
            - Provide a detailed description of the company's business model, industry, and positioning
            - Provide a detailed description of the company's financial performance
            - Provide a detailed description of the company's strategic rationale and growth drivers
            - Provide a detailed description of the company's risk assessment and mitigation strategies
            - Provide a detailed description of the company's recommendation and next steps
        3. Ask
            -  Clearly state the specific approval being sought from the Investment Committee.
        4. Conclusion:
            - Conclude with a formal statement confirming the committee's approval. This should be a direct and affirmative sentence. For example: "The Investment Committee approves the transaction to acquire Secab through the Numeris platform, based on the terms and financial structure presented."

        ## Investment memo to be summarized:
        {ic_text}
        ## Constrictions
        - Follow the Output format
        - The IC minutes must be less than 2 pages long
        '''
        model = model_name
        genai.configure(api_key=os.environ["GOOGLE_API_KEY"])
        gmodel = genai.GenerativeModel(model)
        prompt = Prompt_Template.format(ic_text=text)
        response = gmodel.generate_content(prompt)
        out=response.text
        return out
    def output_docx(self, out: str, path: str = None, title: str = "Investment Committee Minutes"):
        """
        Create a nicely formatted .docx file from the IC minutes text.

        Args:
            out: The generated IC minutes text (ideally containing headings like '1. Context', etc.).
            path: Optional output path (e.g., 'IC_minutes.docx'). If None, a timestamped file is created.
            title: Optional title for the document header.

        Returns:
            The path to the saved .docx file.
        """
        # ---- helpers -------------------------------------------------------
        def _add_page_number_footer(section):
            """
            Adds 'Page X of Y' to the footer using Word fields.
            """
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Build: "Page " + PAGE field + " of " + NUMPAGES field
            run = paragraph.add_run("Page ")
            # PAGE field
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText1 = OxmlElement('w:instrText')
            instrText1.set(qn('xml:space'), 'preserve')
            instrText1.text = "PAGE"

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')

            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')

            r_element = run._r
            r_element.append(fldChar1)
            r_element.append(instrText1)
            r_element.append(fldChar2)
            r_element.append(fldChar3)

            paragraph.add_run(" of ")

            # NUMPAGES field
            run2 = paragraph.add_run()
            fldChar1b = OxmlElement('w:fldChar')
            fldChar1b.set(qn('w:fldCharType'), 'begin')

            instrText2 = OxmlElement('w:instrText')
            instrText2.set(qn('xml:space'), 'preserve')
            instrText2.text = "NUMPAGES"

            fldChar2b = OxmlElement('w:fldChar')
            fldChar2b.set(qn('w:fldCharType'), 'separate')

            fldChar3b = OxmlElement('w:fldChar')
            fldChar3b.set(qn('w:fldCharType'), 'end')

            r2 = run2._r
            r2.append(fldChar1b)
            r2.append(instrText2)
            r2.append(fldChar2b)
            r2.append(fldChar3b)

        def _set_normal_font(doc, name="Calibri", size_pt=11):
            style = doc.styles['Normal']
            style.font.name = name
            style.font.size = Pt(size_pt)

        def _is_top_heading(line: str) -> bool:
            # Matches "1. Context", "2. Investment Overview", "3. Ask", "4. Conclusion"
            return bool(re.match(r"^\s*(\d+)\.\s+.+$", line)) or bool(re.match(r"^#{1,3}\s+.+$", line))

        def _is_bullet(line: str) -> bool:
            return line.strip().startswith("- ")

        def _is_numbered_item(line: str) -> bool:
            # Matches "1) text", "1. text", "(1) text"
            return bool(re.match(r"^\s*(\(?\d+\)?[.)])\s+.+$", line.strip()))

        def _clean_heading_text(line: str) -> str:
            # Remove leading "##", "#" markdown or "1. " numbering
            line = re.sub(r"^\s*#{1,6}\s+", "", line).strip()
            line = re.sub(r"^\s*\d+\.\s+", "", line).strip()
            return line

        # ---- create doc & setup -------------------------------------------
        doc = Document()
        _set_normal_font(doc, name="Calibri", size_pt=11)

        # Page margins (Word default is OK; tune if needed)
        for section in doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Title
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date line
        date_p = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # spacer

        # ---- parse & write content ----------------------------------------
        lines = out.splitlines()

        # If the text is in a single block, create lines by splitting on double newlines
        if len(lines) <= 1 and "\n\n" in out:
            lines = []
            for block in out.split("\n\n"):
                lines.extend(block.splitlines())
                lines.append("")  # preserve paragraph breaks

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if _is_top_heading(line):
                # Major section heading (Context, Investment Overview, etc.)
                clean_text = _clean_heading_text(line)
                heading = doc.add_heading(clean_text, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            elif _is_bullet(line):
                # Bullet point
                p = doc.add_paragraph(line[2:].strip(), style='List Bullet')

            elif _is_numbered_item(line):
                # Numbered list item
                # Extract the text after the number/marker
                match = re.match(r"^\s*(\(?\d+\)?[.)])\s+(.+)$", line.strip())
                if match:
                    text_part = match.group(2)
                    p = doc.add_paragraph(text_part, style='List Number')
                else:
                    # Fallback: just add as normal paragraph
                    p = doc.add_paragraph(line)

            else:
                # Regular paragraph
                p = doc.add_paragraph(line)

        # Add page numbers
        for section in doc.sections:
            _add_page_number_footer(section)

        # ---- save ---------------------------------------------------------
        if path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            path = f"IC_Minutes_{timestamp}.docx"

        doc.save(path)
        print(f"Document saved to: {path}")
        return path
