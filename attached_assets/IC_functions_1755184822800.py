import fitz  # PyMuPDF
import google.generativeai as genai
from PIL import Image
import io
import json
import os
from typing import List, Dict, Any
import time
from dotenv import load_dotenv
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import re

##############
##%

class ICMemoAnalyzer:
    def __init__(self, api_key: str, model_name: str = "gemini-2.5-flash-lite"):
        """
        Initialize the IC Memo Analyzer with Gemini Vision

        Args:
            api_key: Your Google AI API key
            model_name: Gemini model to use (gemini-1.5-pro recommended for vision)
        """
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel(model_name)

    def pdf_to_images(self, pdf_path: str, dpi: int = 300) -> List[Image.Image]:
        """
        Convert PDF pages to high-resolution images

        Args:
            pdf_path: Path to the PDF file
            dpi: Resolution for image conversion (higher = better quality)

        Returns:
            List of PIL Images, one per page
        """
        pdf_document = fitz.open(pdf_path)
        images = []

        # Calculate scaling factor for desired DPI
        # PyMuPDF default is 72 DPI
        scale = dpi / 72.0
        matrix = fitz.Matrix(scale, scale)

        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]

            # Convert page to image
            pix = page.get_pixmap(matrix=matrix)
            img_data = pix.tobytes("png")

            # Convert to PIL Image
            image = Image.open(io.BytesIO(img_data))
            images.append(image)

            print(f"Converted page {page_num + 1}/{pdf_document.page_count}")

        pdf_document.close()
        return images

    def analyze_page(self, image: Image.Image, page_num: int) -> Dict[str, Any]:
        """
        Analyze a single page using Gemini Vision

        Args:
            image: PIL Image of the page
            page_num: Page number for reference

        Returns:
            Dictionary containing extracted information
        """
        prompt = """
        # Persona
        You are a financial analyst working at Strada Partners, a  a small-cap private equity firm based in Belgium.
        # Task
        Your task is to analyze IC (Investment Committee) memos and extract key information.
        Be thorough and precise with numbers, percentages, and specific details shown in the document.
        # Output
        Output a text with all the relevant information, such that it can be further processed
        """

        try:
            response = self.model.generate_content([prompt, image])

            return {
                "page_number": page_num + 1,
                "content": response.text,
                "status": "success"
            }

        except Exception as e:
            return {
                "page_number": page_num + 1,
                "content": f"Error analyzing page: {str(e)}",
                "status": "error"
            }

    def analyze_pdf(self, pdf_path: str, output_file: str = None, delay: float = 1.0) -> List[Dict[str, Any]]:
        """
        Analyze entire PDF document

        Args:
            pdf_path: Path to the PDF file
            output_file: Optional path to save results as JSON
            delay: Delay between API calls to avoid rate limits

        Returns:
            List of analysis results, one per page
        """
        print(f"Starting analysis of {pdf_path}")

        # Convert PDF to images
        images = self.pdf_to_images(pdf_path)

        # Analyze each page
        results = []
        for i, image in enumerate(images):
            print(f"Analyzing page {i + 1}/{len(images)}...")

            result = self.analyze_page(image, i)
            results.append(result)

            # Add delay to avoid rate limiting
            if i < len(images) - 1:  # Don't delay after the last page
                time.sleep(delay)

        # Save results if output file specified
        if output_file:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
            print(f"Results saved to {output_file}")

        return results

    def create_IC_minutes(self,results, model_name="gemini-2.5-pro"):
        '''
        this function creates IC minutes from the results of the analysis
        :param results:
        :param model_name:
        :return:
        '''
        text = ""
        for i in range(len(results)):
            text += "results from page " + str(i + 1) + "\n" + results[i]["content"] + "\n\n"


        Prompt_Template = '''
        ## Persona
        You are an expert investment analyst tasked with summarizing a detailed investment memo into a formal and concise Investment Committee (IC) Minutes document. 
        You work at Strada Partners, a small-cap private equity firm based in Belgium. You need to draft IC minutes for the investment memo below:
        ## Task
        Your task is to analyze IC (Investment Committee) memos and extract key information.
        Be thorough and precise with numbers, percentages, and specific details shown in the document.
        ## Input
        The input is a PDF file containing an investment memo. The PDF file is provided as a single file.
        ## Output
        Your output must be structured into the following four sections, in this specific order:
        1. Context
        2. Investment Overview
        3. Ask
        4. Conclusion
        Use the provided text, which is an extraction from an investment memo on "Project Calc," to generate the IC minutes.

        ## Detailed description
        1. Context
            - Introduce the Company being acquired
            - Condense the most relevant information from the memo
            - Provide a brief summary of the company's business model, industry, and positioning
        2. Investment Overview
            - Provide a detailed overview of the investment, including the company's name, the type of investment, the investment objective, and the investment strategy
            - Provide a detailed description of the investment objective and strategy
            - Provide a detailed description of the company's business model, industry, and positioning
            - Provide a detailed description of the company's financial performance
            - Provide a detailed description of the company's strategic rationale and growth drivers
            - Provide a detailed description of the company's risk assessment and mitigation strategies
            - Provide a detailed description of the company's recommendation and next steps
        3. Ask
            -  Clearly state the specific approval being sought from the Investment Committee.
        4. Conclusion:
            - Conclude with a formal statement confirming the committee's approval. This should be a direct and affirmative sentence. For example: "The Investment Committee approves the transaction to acquire Secab through the Numeris platform, based on the terms and financial structure presented."

        ## Investment memo to be summarized:
        {ic_text}
        ## Constrictions
        - Follow the Output format
        - The IC minutes must be less than 2 pages long
        '''
        model = model_name
        genai.configure(api_key=os.environ["GOOGLE_API_KEY"])
        gmodel = genai.GenerativeModel(model)
        prompt = Prompt_Template.format(ic_text=text)
        response = gmodel.generate_content(prompt)
        out=response.text
        return out
    def output_docx(self, out: str, path: str = None, title: str = "Investment Committee Minutes"):
        """
        Create a nicely formatted .docx file from the IC minutes text.

        Args:
            out: The generated IC minutes text (ideally containing headings like '1. Context', etc.).
            path: Optional output path (e.g., 'IC_minutes.docx'). If None, a timestamped file is created.
            title: Optional title for the document header.

        Returns:
            The path to the saved .docx file.
        """
        # ---- helpers -------------------------------------------------------
        def _add_page_number_footer(section):
            """
            Adds 'Page X of Y' to the footer using Word fields.
            """
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Build: "Page " + PAGE field + " of " + NUMPAGES field
            run = paragraph.add_run("Page ")
            # PAGE field
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText1 = OxmlElement('w:instrText')
            instrText1.set(qn('xml:space'), 'preserve')
            instrText1.text = "PAGE"

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')

            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')

            r_element = run._r
            r_element.append(fldChar1)
            r_element.append(instrText1)
            r_element.append(fldChar2)
            r_element.append(fldChar3)

            paragraph.add_run(" of ")

            # NUMPAGES field
            run2 = paragraph.add_run()
            fldChar1b = OxmlElement('w:fldChar')
            fldChar1b.set(qn('w:fldCharType'), 'begin')

            instrText2 = OxmlElement('w:instrText')
            instrText2.set(qn('xml:space'), 'preserve')
            instrText2.text = "NUMPAGES"

            fldChar2b = OxmlElement('w:fldChar')
            fldChar2b.set(qn('w:fldCharType'), 'separate')

            fldChar3b = OxmlElement('w:fldChar')
            fldChar3b.set(qn('w:fldCharType'), 'end')

            r2 = run2._r
            r2.append(fldChar1b)
            r2.append(instrText2)
            r2.append(fldChar2b)
            r2.append(fldChar3b)

        def _set_normal_font(doc, name="Calibri", size_pt=11):
            style = doc.styles['Normal']
            style.font.name = name
            style.font.size = Pt(size_pt)

        def _is_top_heading(line: str) -> bool:
            # Matches "1. Context", "2. Investment Overview", "3. Ask", "4. Conclusion"
            return bool(re.match(r"^\s*(\d+)\.\s+.+$", line)) or bool(re.match(r"^#{1,3}\s+.+$", line))

        def _is_bullet(line: str) -> bool:
            return line.strip().startswith("- ")

        def _is_numbered_item(line: str) -> bool:
            # Matches "1) text", "1. text", "(1) text"
            return bool(re.match(r"^\s*(\(?\d+\)?[.)])\s+.+$", line.strip()))

        def _clean_heading_text(line: str) -> str:
            # Remove leading "##", "#" markdown or "1. " numbering
            line = re.sub(r"^\s*#{1,6}\s+", "", line).strip()
            line = re.sub(r"^\s*\d+\.\s+", "", line).strip()
            return line

        # ---- create doc & setup -------------------------------------------
        doc = Document()
        _set_normal_font(doc, name="Calibri", size_pt=11)

        # Page margins (Word default is OK; tune if needed)
        for section in doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Title
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date line
        date_p = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # spacer

        # ---- parse & write content ----------------------------------------
        lines = out.splitlines()

        # If the text is in a single block, create lines by splitting on double newlines
        if len(lines) <= 1 and "\n\n" in out:
            lines = []
            for block in out.split("\n\n"):
                lines.extend(block.splitlines())

        # Track whether we are inside a list to apply consistent styles
        for raw_line in lines:
            line = raw_line.rstrip()

            if not line.strip():
                # Blank line -> add a small spacer paragraph
                doc.add_paragraph()
                continue

            if _is_top_heading(line):
                # Top-level section heading
                heading_text = _clean_heading_text(line)
                doc.add_heading(heading_text, level=1)
                continue

            if _is_bullet(line):
                # Bullet list
                p = doc.add_paragraph(line.strip()[2:])
                p.style = "List Bullet"
                continue

            if _is_numbered_item(line):
                # Numbered list
                # Remove the leading 1., 1) or (1)
                text = re.sub(r"^\s*(\(?\d+\)?[.)])\s+", "", line.strip())
                p = doc.add_paragraph(text)
                p.style = "List Number"
                continue

            # Fallback: normal paragraph
            doc.add_paragraph(line)

        # Footer with page numbering
        for section in doc.sections:
            _add_page_number_footer(section)

        # ---- save ----------------------------------------------------------
        if path is None:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            path = f"IC_minutes_{ts}.docx"

        # Ensure directory exists if a nested path is given
        os.makedirs(os.path.dirname(path), exist_ok=True) if os.path.dirname(path) else None
        doc.save(path)
        return path




####
#%##
# import fitz  # PyMuPDF
# import google.generativeai as genai
# from PIL import Image
# import io
# import json
# import os
# from typing import List, Dict, Any, Optional
# import time
# import concurrent.futures
# from threading import Lock
# import re
#
# class FastICMemoAnalyzer:
#     def __init__(self, api_key: str,
#                  text_model: str = "gemini-2.5-flash",
#                  vision_model: str = "gemini-2.5-flash"):
#         """
#         Fast analyzer that uses text extraction first, vision only when needed
#
#         Args:
#             api_key: Your Google AI API key
#             text_model: Gemini model for text analysis
#             vision_model: Gemini model for vision analysis
#         """
#         genai.configure(api_key=api_key)
#         self.text_model = genai.GenerativeModel(text_model)
#         self.vision_model = genai.GenerativeModel(vision_model)
#         self.api_lock = Lock()  # Prevent concurrent API calls
#
#     def extract_text_and_metadata(self, pdf_path: str) -> Dict[str, Any]:
#         """
#         Extract text and identify pages that might need vision analysis
#
#         Returns:
#             Dictionary with text content and metadata for each page
#         """
#         pdf_document = fitz.open(pdf_path)
#         pages_data = []
#
#         for page_num in range(pdf_document.page_count):
#             page = pdf_document[page_num]
#
#             # Extract text
#             text = page.get_text()
#
#             # Get page metadata
#             images = page.get_images()
#             drawings = page.get_drawings()
#
#             # Determine if vision analysis might be needed
#             needs_vision = self._needs_vision_analysis(text, images, drawings)
#
#             pages_data.append({
#                 "page_number": page_num + 1,
#                 "text": text.strip(),
#                 "text_length": len(text.strip()),
#                 "image_count": len(images),
#                 "drawing_count": len(drawings),
#                 "needs_vision": needs_vision,
#                 "page_obj": page if needs_vision else None
#             })
#
#             print(f"Processed page {page_num + 1}/{pdf_document.page_count} - "
#                   f"Text: {len(text)} chars, Images: {len(images)}, "
#                   f"Vision needed: {needs_vision}")
#
#         # Don't close document yet if we need vision analysis
#         if not any(p["needs_vision"] for p in pages_data):
#             pdf_document.close()
#
#         return {
#             "pages": pages_data,
#             "pdf_document": pdf_document if any(p["needs_vision"] for p in pages_data) else None,
#             "total_pages": len(pages_data)
#         }
#
#     def _needs_vision_analysis(self, text: str, images: list, drawings: list) -> bool:
#         """
#         Determine if a page needs vision analysis based on content
#         """
#         # If text is very short but has images/drawings, might need vision
#         if len(text.strip()) < 100 and (images or drawings):
#             return True
#
#         # Look for indicators of complex visual content
#         visual_indicators = [
#             "chart", "graph", "diagram", "figure", "exhibit",
#             "organizational structure", "flow", "timeline"
#         ]
#
#         text_lower = text.lower()
#         if any(indicator in text_lower for indicator in visual_indicators) and images:
#             return True
#
#         # If there are many images, might be visual-heavy
#         if len(images) > 3:
#             return True
#
#         return False
#
#     def analyze_text_batch(self, pages_batch: List[Dict], batch_size: int = 10) -> List[Dict]:
#         """
#         Analyze a batch of text-based pages efficiently
#         """
#         # Combine multiple pages into one prompt for efficiency
#         combined_text = ""
#         page_markers = []
#
#         for page_data in pages_batch:
#             page_num = page_data["page_number"]
#             text = page_data["text"]
#
#             if text:  # Only include pages with text
#                 marker = f"\n\n===== PAGE {page_num} =====\n"
#                 combined_text += marker + text
#                 page_markers.append(page_num)
#
#         if not combined_text.strip():
#             return [{"page_number": p["page_number"], "content": "No text content",
#                     "status": "empty"} for p in pages_batch]
#
#         prompt = f"""
#         # Persona
#         You are a financial analyst working at Strada Partners, a  a small-cap private equity firm based in Belgium.
#         # Task
#         Your task is to analyze IC (Investment Committee) memos and extract key information.
#         Be thorough and precise with numbers, percentages, and specific details shown in the document.
#         # Output
#         Output a text with all the relevant information, such that it can be further processed
#
#         Content to analyze:
#         {combined_text}
#
#         Please structure your response clearly by page number
#         """
#
#         try:
#             with self.api_lock:
#                 response = self.text_model.generate_content(prompt)
#
#             return [{
#                 "page_numbers": page_markers,
#                 "content": response.text,
#                 "status": "success",
#                 "method": "batch_text_analysis"
#             }]
#
#         except Exception as e:
#             return [{
#                 "page_numbers": page_markers,
#                 "content": f"Error in batch analysis: {str(e)}",
#                 "status": "error",
#                 "method": "batch_text_analysis"
#             }]
#
#     def analyze_vision_page(self, page_obj, page_num: int, dpi: int = 200) -> Dict:
#         """
#         Analyze a single page using vision model (lower DPI for speed)
#         """
#         try:
#             # Convert to image with lower DPI for speed
#             scale = dpi / 72.0
#             matrix = fitz.Matrix(scale, scale)
#             pix = page_obj.get_pixmap(matrix=matrix)
#             img_data = pix.tobytes("png")
#             image = Image.open(io.BytesIO(img_data))
#
#             prompt = """
#             # Persona
#         You are a financial analyst working at Strada Partners, a  a small-cap private equity firm based in Belgium.
#         # Task
#         Your task is to analyze IC (Investment Committee) memos and extract key information.
#             This page likely contains charts, diagrams, or visual elements that weren't captured in text extraction.
#
#             Focus on:
#             - Data from any charts, graphs, or tables
#             - Organizational structures or process flows
#             - Visual relationships between entities
#             - Any financial data presented visually
#             - Key insights that require visual context
#
#             Be concise but thorough about visual elements.
#             """
#
#             with self.api_lock:
#                 response = self.vision_model.generate_content([prompt, image])
#
#             return {
#                 "page_number": page_num,
#                 "content": response.text,
#                 "status": "success",
#                 "method": "vision_analysis"
#             }
#
#         except Exception as e:
#             return {
#                 "page_number": page_num,
#                 "content": f"Vision analysis error: {str(e)}",
#                 "status": "error",
#                 "method": "vision_analysis"
#             }
#
#     def analyze_pdf_fast(self, pdf_path: str, batch_size: int = 10,
#                         max_workers: int = 3) -> Dict[str, Any]:
#         """
#         Fast analysis of large PDFs using hybrid approach
#
#         Args:
#             pdf_path: Path to PDF
#             batch_size: Number of pages to analyze together
#             max_workers: Number of concurrent threads (be careful with API limits)
#         """
#         print(f"Starting fast analysis of {pdf_path}")
#         start_time = time.time()
#
#         # Extract all text and metadata
#         print("Extracting text and analyzing structure...")
#         pdf_data = self.extract_text_and_metadata(pdf_path)
#         pages = pdf_data["pages"]
#
#         # Separate text-only and vision-needed pages
#         text_pages = [p for p in pages if not p["needs_vision"] and p["text"]]
#         vision_pages = [p for p in pages if p["needs_vision"]]
#
#         print(f"Analysis plan:")
#         print(f"  - {len(text_pages)} pages for fast text analysis")
#         print(f"  - {len(vision_pages)} pages need vision analysis")
#         print(f"  - Using batch size of {batch_size} for text analysis")
#
#         all_results = []
#
#         # Process text pages in batches (much faster)
#         if text_pages:
#             print("\nProcessing text-based pages in batches...")
#             for i in range(0, len(text_pages), batch_size):
#                 batch = text_pages[i:i + batch_size]
#                 batch_num = i // batch_size + 1
#                 total_batches = (len(text_pages) + batch_size - 1) // batch_size
#
#                 print(f"  Processing text batch {batch_num}/{total_batches} "
#                       f"(pages {batch[0]['page_number']}-{batch[-1]['page_number']})")
#
#                 batch_results = self.analyze_text_batch(batch, batch_size)
#                 all_results.extend(batch_results)
#
#                 # Small delay between batches
#                 time.sleep(0.5)
#
#         # Process vision pages (slower, but only when needed)
#         if vision_pages:
#             print(f"\nProcessing {len(vision_pages)} pages with vision analysis...")
#
#             # Use threading for vision pages (be careful with rate limits)
#             def process_vision_page(page_data):
#                 return self.analyze_vision_page(page_data["page_obj"], page_data["page_number"])
#
#             with concurrent.futures.ThreadPoolExecutor(max_workers=min(max_workers, len(vision_pages))) as executor:
#                 vision_results = list(executor.map(process_vision_page, vision_pages))
#                 all_results.extend(vision_results)
#
#         # Clean up
#         if pdf_data["pdf_document"]:
#             pdf_data["pdf_document"].close()
#
#         elapsed_time = time.time() - start_time
#
#         analysis_summary = {
#             "total_pages": len(pages),
#             "text_pages": len(text_pages),
#             "vision_pages": len(vision_pages),
#             "processing_time": elapsed_time,
#             "pages_per_minute": len(pages) / (elapsed_time / 60) if elapsed_time > 0 else 0
#         }
#
#         print(f"\nAnalysis complete!")
#         print(f"  Total time: {elapsed_time:.1f} seconds")
#         print(f"  Speed: {analysis_summary['pages_per_minute']:.1f} pages/minute")
#
#         return {
#             "results": all_results,
#             "summary": analysis_summary,
#             "pdf_path": pdf_path
#         }
#
#     def create_executive_summary(self, analysis_results: Dict[str, Any]) -> str:
#         """
#         Create executive summary from fast analysis results
#         """
#         # Combine all analysis content
#         all_content = []
#         for result in analysis_results["results"]:
#             if result["status"] == "success":
#                 if "page_numbers" in result:  # Batch result
#                     all_content.append(f"Pages {result['page_numbers']}: {result['content']}")
#                 else:  # Single page result
#                     all_content.append(f"Page {result['page_number']}: {result['content']}")
#
#         combined_content = "\n\n".join(all_content)
#
#         summary_prompt = f"""
#         Create a comprehensive executive summary from this IC memo analysis:
#
#         {combined_content}
#
#         Structure the summary as:
#         1. **Executive Summary** - Key investment decision and rationale
#         2. **Financial Highlights** - Critical numbers and performance metrics
#         3. **Strategic Assessment** - Market position and growth opportunities
#         4. **Risk Analysis** - Major risks and mitigation strategies
#         5. **Recommendation** - Clear action items and timeline
#
#         Focus on actionable insights and key decision points for the investment committee.
#         """
#
#         try:
#             response = self.text_model.generate_content(summary_prompt)
#             return response.text
#         except Exception as e:
#             return f"Error creating summary: {str(e)}"
